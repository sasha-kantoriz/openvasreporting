# -*- coding: utf-8 -*-
#
#
# Project name: OpenVAS Reporting: A tool to convert OpenVAS XML reports into Excel files.
# Project URL: https://github.com/TheGroundZero/openvasreporting

import re

from collections import Counter

from .config import Config
from .parsed_data import ResultTree, Host, Vulnerability

# DEBUG
#import sys
#import logging
#logging.basicConfig(stream=sys.stderr, level=logging.WARNING,
#                     format="%(asctime)s | %(levelname)s | %(name)s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S")


def implemented_exporters():
    """
    Enum-link instance containing references to already implemented exporter function

    > implemented_exporters()[key](param[s])
    
    key is a concatenation of the report-type arg + '-' + format arg

    :return: Pointer to exporter function
    """
    return {
        'vulnerability-xlsx': export_to_excel_by_vuln,
        'vulnerability-docx': export_to_word_by_vuln,
        'vulnerability-csv': export_to_csv_by_vuln,
        'host-docx': export_to_word_by_host,
        'host-xlsx': export_to_excel_by_host,
        'host-csv': export_to_csv_by_host,
        'summary-csv': export_summary_to_csv
    }

def _get_collections(vuln_info):
    """
    Sort vulnerability list info according to CVSS (desc) and Name (asc).
    Provide collections to be used in export.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :return: vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family
    :rtype vuln_info: list(Vulnerability)
    :rtype vuln_levels: Counter
    :rtype vuln_host_by_level: Counter
    :rtype vuln_by_family: Counter
    """
    vuln_info.sort(key=lambda key: key.name)
    vuln_info.sort(key=lambda key: key.cvss, reverse=True)
    vuln_levels = Counter()
    vuln_host_by_level = Counter()
    vuln_by_family = Counter()
    # collect host names
    vuln_hostcount_by_level =[[] for _ in range(5)]
    level_choices = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3, 'none': 4}

    for i, vuln in enumerate(vuln_info, 1):
        vuln_levels[vuln.level.lower()] += 1
        # add host names to list so we count unquie hosts per level
        level_index = level_choices.get(vuln.level.lower())

        for i, (host, port) in enumerate(vuln.hosts, 1):    
            if host.ip not in vuln_hostcount_by_level[level_index]:
                vuln_hostcount_by_level[level_index].append(host.ip)       

        vuln_by_family[vuln.family] += 1

    # now count hosts per level and return
    for level in Config.levels().values():
        vuln_host_by_level[level] = len((vuln_hostcount_by_level[level_choices.get(level.lower())]))

    return vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family


def export_to_excel_by_vuln(vuln_info, template=None, output_file='openvas_report.xlsx'):
    """
    Export vulnerabilities info in an Excel file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)
    :param template: Not supported in xlsx-output
    :type template: NoneType

    :param output_file: Filename of the Excel file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import xlsxwriter

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    # if template is not None:
    #     raise NotImplementedError("Use of template is not supported in XSLX-output.")

    vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family = _get_collections(vuln_info)

    # ====================
    # FUNCTIONS
    # ====================
    def __row_height(text, width):
        return (max((len(text) // width), text.count('\n')) + 1) * 15

    workbook = xlsxwriter.Workbook(output_file)

    workbook.set_properties({
        'title': output_file,
        'subject': 'OpenVAS report',
        'author': 'TheGroundZero',
        'category': 'report',
        'keywords': 'OpenVAS, report',
        'comments': 'TheGroundZero (https://github.com/TheGroundZero)'})

    # ====================
    # FORMATTING
    # ====================
    workbook.formats[0].set_font_name('Tahoma')

    format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'center', 'valign': 'vcenter', 'border': 1})
    format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                               'font_color': 'white', 'bold': True,
                                               'align': 'center', 'valign': 'vcenter',
                                               'border': 1,
                                               'bg_color': Config.colors()['blue']})
    format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                              'align': 'left', 'valign': 'top',
                                              'border': 1, 'text_wrap': 1})
    format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top'})
    format_align_border = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_toc = {
        'critical': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                         'align': 'center', 'valign': 'top',
                                         'border': 1,
                                         'bg_color': Config.colors()['critical']}),
        'high': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['high']}),
        'medium': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                       'align': 'center', 'valign': 'top',
                                       'border': 1, 'bg_color': Config.colors()['medium']}),
        'low': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                    'align': 'center', 'valign': 'top',
                                    'border': 1, 'bg_color': Config.colors()['low']}),
        'none': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['none']})
    }

    # ====================
    # SUMMARY SHEET
    # ====================
    sheet_name = "Summary"
    ws_sum = workbook.add_worksheet(sheet_name)
    ws_sum.set_tab_color(Config.colors()['blue'])

    ws_sum.set_column("A:A", 7, format_align_center)
    ws_sum.set_column("B:B", 25, format_align_center)
    ws_sum.set_column("C:C", 24, format_align_center)
    ws_sum.set_column("D:D", 20, format_align_center)
    ws_sum.set_column("E:E", 7, format_align_center)

    # --------------------
    # VULN SUMMARY
    # --------------------
    ws_sum.merge_range("B2:D2", "VULNERABILITY SUMMARY", format_sheet_title_content)
    ws_sum.write("B3", "Threat Level", format_table_titles)
    ws_sum.write("C3", "Vulnerabilities", format_table_titles)
    ws_sum.write("D3", "Affected hosts", format_table_titles)

    for i, level in enumerate(Config.levels().values(), 4):
        ws_sum.write("B{}".format(i), level.capitalize(), format_sheet_title_content)
        ws_sum.write("C{}".format(i), vuln_levels[level], format_align_border)
        ws_sum.write("D{}".format(i), vuln_host_by_level[level], format_align_border)

    ws_sum.write("B9", "Total", format_table_titles)
    ws_sum.write_formula("C9", "=SUM($C$4:$C$8)", format_table_titles)
    ws_sum.write_formula("D9", "=SUM($D$4:$D$8)", format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_summary = workbook.add_chart({'type': 'pie'})
    chart_vulns_summary.add_series({
        'name': 'vulnerability summary by affected hosts',
        'categories': '={}!B4:B8'.format(sheet_name),
        'values': '={}!D4:D8'.format(sheet_name),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
        'points': [
            {'fill': {'color': Config.colors()['critical']}},
            {'fill': {'color': Config.colors()['high']}},
            {'fill': {'color': Config.colors()['medium']}},
            {'fill': {'color': Config.colors()['low']}},
            {'fill': {'color': Config.colors()['none']}},
        ],
    })
    chart_vulns_summary.set_title({'name': 'Vulnerability summary', 'overlay': False, 'name_font': {'name': 'Tahoma'}})
    chart_vulns_summary.set_size({'width': 500, 'height': 300})
    chart_vulns_summary.set_legend({'position': 'right', 'font': {'name': 'Tahoma'}})
    ws_sum.insert_chart("F2", chart_vulns_summary)

    # --------------------
    # VULN BY FAMILY
    # --------------------
    ws_sum.merge_range("B19:C19", "VULNERABILITIES BY FAMILY", format_sheet_title_content)
    ws_sum.write("B20", "Family", format_table_titles)
    ws_sum.write("C20", "Vulnerabilities", format_table_titles)

    last = 21
    for i, (family, number) in enumerate(iter(vuln_by_family.items()), last):
        ws_sum.write("B{}".format(i), family, format_align_border)
        ws_sum.write("C{}".format(i), number, format_align_border)
        last = i

    ws_sum.write("B{}".format(str(last + 1)), "Total", format_table_titles)
    ws_sum.write_formula("C{}".format(str(last + 1)), "=SUM($C$21:$C${})".format(last), format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_by_family = workbook.add_chart({'type': 'pie'})
    chart_vulns_by_family.add_series({
        'name': 'vulnerability summary by family',
        'categories': '={}!B21:B{}'.format(sheet_name, last),
        'values': '={}!C21:C{}'.format(sheet_name, last),
        'data_labels': {'value': True, 'position': 'best_fit', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
    })
    chart_vulns_by_family.set_title({'name': 'Vulnerabilities by family', 'overlay': False,
                                     'name_font': {'name': 'Tahoma'}})
    chart_vulns_by_family.set_size({'width': 500, 'height': 500})
    chart_vulns_by_family.set_legend({'position': 'bottom', 'font': {'name': 'Tahoma'}})
    ws_sum.insert_chart("F19", chart_vulns_by_family)

    # ====================
    # TABLE OF CONTENTS
    # ====================
    sheet_name = "TOC"
    ws_toc = workbook.add_worksheet(sheet_name)
    ws_toc.set_tab_color(Config.colors()['blue'])

    ws_toc.set_column("A:A", 7)
    ws_toc.set_column("B:B", 5)
    ws_toc.set_column("C:C", 70)
    ws_toc.set_column("D:D", 15)
    ws_toc.set_column("E:E", 50)
    ws_toc.set_column("F:F", 7)

    ws_toc.merge_range("B2:E2", "TABLE OF CONTENTS", format_sheet_title_content)
    ws_toc.write("B3", "No.", format_table_titles)
    ws_toc.write("C3", "Vulnerability", format_table_titles)
    ws_toc.write("D3", "CVSS Score", format_table_titles)
    ws_toc.write("E3", "Hosts", format_table_titles)

    # ====================
    # VULN SHEETS
    # ====================
    for i, vuln in enumerate(vuln_info, 1):
        name = re.sub(r"[\[\]\\\'\"&@#():*?/]", "", vuln.name)
        if len(name) > 27:
            name = "{}..{}".format(name[0:15], name[-10:])
        name = "{:03X}_{}".format(i, name)
        ws_vuln = workbook.add_worksheet(name)
        ws_vuln.set_tab_color(Config.colors()[vuln.level.lower()])

        # --------------------
        # TABLE OF CONTENTS
        # --------------------
        ws_toc.write("B{}".format(i + 3), "{:03X}".format(i), format_table_cells)
        ws_toc.write_url("C{}".format(i + 3), "internal:'{}'!A1".format(name), format_table_cells, string=vuln.name)
        ws_toc.write("D{}".format(i + 3), "{:.1f} ({})".format(vuln.cvss, vuln.level.capitalize()),
                     format_toc[vuln.level])
        ws_toc.write("E{}".format(i + 3), "{}".format(', '.join([host.ip for host, _ in vuln.hosts])),
                     format_table_cells)
        ws_vuln.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), i + 3), format_align_center,
                          string="<< TOC")
        ws_toc.set_row(i + 3, __row_height(name, 150), None)

        # --------------------
        # VULN INFO
        # --------------------
        ws_vuln.set_column("A:A", 7, format_align_center)
        ws_vuln.set_column("B:B", 20, format_align_center)
        ws_vuln.set_column("C:C", 20, format_align_center)
        ws_vuln.set_column("D:D", 50, format_align_center)
        ws_vuln.set_column("E:E", 15, format_align_center)
        ws_vuln.set_column("F:F", 15, format_align_center)
        ws_vuln.set_column("G:G", 20, format_align_center)
        ws_vuln.set_column("H:H", 7, format_align_center)
        content_width = 120

        ws_vuln.write('B2', "Title", format_table_titles)
        ws_vuln.merge_range("C2:G2", vuln.name, format_sheet_title_content)
        ws_vuln.set_row(1, __row_height(vuln.name, content_width), None)

        ws_vuln.write('B3', "Description", format_table_titles)
        ws_vuln.merge_range("C3:G3", vuln.description, format_table_cells)
        ws_vuln.set_row(2, __row_height(vuln.description, content_width), None)

        ws_vuln.write('B4', "Impact", format_table_titles)
        ws_vuln.merge_range("C4:G4", vuln.impact, format_table_cells)
        ws_vuln.set_row(3, __row_height(vuln.impact, content_width), None)

        ws_vuln.write('B5', "Recommendation", format_table_titles)
        ws_vuln.merge_range("C5:G5", vuln.solution, format_table_cells)
        ws_vuln.set_row(4, __row_height(vuln.solution, content_width), None)

        ws_vuln.write('B6', "Details", format_table_titles)
        ws_vuln.merge_range("C6:G6", vuln.insight, format_table_cells)
        ws_vuln.set_row(5, __row_height(vuln.insight, content_width), None)

        ws_vuln.write('B7', "CVEs", format_table_titles)
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"
        ws_vuln.merge_range("C7:G7", cves, format_table_cells)
        ws_vuln.set_row(6, __row_height(cves, content_width), None)

        ws_vuln.write('B8', "CVSS", format_table_titles)
        cvss = float(vuln.cvss)
        if cvss >= 0.0:
            ws_vuln.merge_range("C8:G8", "{:.1f}".format(cvss), format_table_cells)
        else:
            ws_vuln.merge_range("C8:G8", "{}".format("No CVSS"), format_table_cells)

        ws_vuln.write('B9', "Level", format_table_titles)
        ws_vuln.merge_range("C9:G9", vuln.level.capitalize(), format_table_cells)

        ws_vuln.write('B10', "Family", format_table_titles)
        ws_vuln.merge_range("C10:G10", vuln.family, format_table_cells)

        ws_vuln.write('B11', "References", format_table_titles)
        ws_vuln.merge_range("C11:G11", " {}".format(vuln.references), format_table_cells)
        ws_vuln.set_row(10, __row_height(vuln.references, content_width), None)

        ws_vuln.write('C13', "IP", format_table_titles)
        ws_vuln.write('D13', "Host name", format_table_titles)
        ws_vuln.write('E13', "Port number", format_table_titles)
        ws_vuln.write('F13', "Port protocol", format_table_titles)
        ws_vuln.write('G13', "Result", format_table_titles)

        # --------------------
        # AFFECTED HOSTS
        # --------------------
        for j, (host, port) in enumerate(vuln.hosts, 14):

            ws_vuln.write("C{}".format(j), host.ip)
            ws_vuln.write("D{}".format(j), host.host_name if host.host_name else "-")

            if port:
                ws_vuln.write("E{}".format(j), "" if port.number == 0 else port.number)
                ws_vuln.write("F{}".format(j), port.protocol)
                ws_vuln.write("G{}".format(j), port.result, format_table_cells)
                ws_vuln.set_row(j, __row_height(port.result, content_width), None)
            else:
                ws_vuln.write("E{}".format(j), "No port info")

    workbook.close()


def export_to_word_by_vuln(vuln_info, template, output_file='openvas_report.docx'):
    """
    Export vulnerabilities info in a Word file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param output_file: Filename of the Excel file
    :type output_file: str
    
    :param template: Path to Docx template
    :type template: str

    :raises: TypeError
    """

    import matplotlib.pyplot as plt
    import numpy as np
    import tempfile
    import os

    from docx import Document
    from docx.oxml.shared import qn, OxmlElement
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        if not isinstance(template, str):
            raise TypeError("Expected str, got '{}' instead".format(type(template)))
    else:
        # == HAMMER PROGRAMMING (beat it into submission) ==
        # I had to use pkg_resources because I couldn't find this template any other way. 
        import pkg_resources
        template = pkg_resources.resource_filename('openvasreporting', 'src/openvas-template.docx')
    
    #mpl_logger = logging.getLogger('matplotlib')
    #mpl_logger.setLevel(logging.NOTSET)
    
    vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family = _get_collections(vuln_info)

    # ====================
    # DOCUMENT PROPERTIES
    # ====================
    document = Document(template)

    doc_prop = document.core_properties
    doc_prop.title = "OpenVAS Report"
    doc_prop.category = "Report"

    document.add_paragraph('OpenVAS Report', style='Title')

    # ====================
    # TABLE OF CONTENTS
    # ====================
    document.add_paragraph('Table of Contents', style='Heading 1')

    par = document.add_paragraph()
    run = par.add_run()
    fld_char = OxmlElement('w:fldChar')  # creates a new element
    fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instr_text.text = r'TOC \h \z \t "OV-H1toc;1;OV-H2toc;2;OV-H3toc;3;OV-Finding;3"'

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:t')
    fld_char3.text = "# Right-click to update field. #"
    fld_char2.append(fld_char3)

    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)

    document.add_page_break()

    # ====================
    # MANAGEMENT SUMMARY
    # ====================
    document.add_paragraph('Management Summary', style='OV-H1toc')
    document.add_paragraph('< TYPE YOUR MANAGEMENT SUMMARY HERE >')
    document.add_page_break()

    # ====================
    # TECHNICAL FINDINGS
    # ====================
    document.add_paragraph('Technical Findings', style='OV-H1toc')
    document.add_paragraph('The section below discusses the technical findings.')

    # --------------------
    # SUMMARY TABLE
    # --------------------
    document.add_paragraph('Summary', style='OV-H2toc')

    colors_sum = []
    labels_sum = []
    vuln_sum = []
    aff_sum = []

    table_summary = document.add_table(rows=1, cols=3)
    hdr_cells = table_summary.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Risk level').bold = True
    hdr_cells[1].paragraphs[0].add_run('Vulns number').bold = True
    hdr_cells[2].paragraphs[0].add_run('Affected hosts').bold = True

    # Provide data to table and charts
    for level in Config.levels().values():
        row_cells = table_summary.add_row().cells
        row_cells[0].text = level.capitalize()
        row_cells[1].text = str(vuln_levels[level])
        row_cells[2].text = str(vuln_host_by_level[level])
        colors_sum.append(Config.colors()[level])
        labels_sum.append(level)
        vuln_sum.append(vuln_levels[level])
        aff_sum.append(vuln_host_by_level[level])

    # --------------------
    # CHART
    # --------------------
    fd, path = tempfile.mkstemp(suffix='.png')

    par_chart = document.add_paragraph()
    run_chart = par_chart.add_run()

    plt.figure()

    pos = np.arange(len(labels_sum))
    width = 0.35

    bars_vuln = plt.bar(pos - width / 2, vuln_sum, width, align='center', label='Vulnerabilities',
                        color=colors_sum, edgecolor='black')
    bars_aff = plt.bar(pos + width / 2, aff_sum, width, align='center', label='Affected hosts',
                       color=colors_sum, edgecolor='black', hatch='//')
    plt.title('Vulnerability summary by risk level')
    plt.subplot().set_xticks(pos)
    plt.subplot().set_xticklabels(labels_sum)
    plt.gca().spines['left'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['bottom'].set_position('zero')
    plt.tick_params(top=False, bottom=True, left=False, right=False,
                    labelleft=False, labelbottom=True)
    plt.subplots_adjust(left=0.0, right=1.0)

    def __label_bars(barcontainer):
        for bar in barcontainer:
            height = bar.get_height()
            plt.gca().text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, str(int(height)),
                           ha='center', color='black', fontsize=9)

    __label_bars(bars_vuln)
    __label_bars(bars_aff)

    plt.legend()

    plt.savefig(path)

    # plt.show()  # DEBUG

    run_chart.add_picture(path, width=Cm(8.0))

    plt.figure()

    values = list(vuln_by_family.values())
    pie, tx, autotexts = plt.pie(values, labels=vuln_by_family.keys(), autopct='')
    plt.title('Vulnerability by family')
    for i, txt in enumerate(autotexts):
        txt.set_text('{}'.format(values[i]))
    plt.axis('equal')

    plt.savefig(path, bbox_inches='tight')  # bbox_inches fixes labels being cut, however only on save not on show

    # plt.show()  # DEBUG

    run_chart.add_picture(path, width=Cm(8.0))
    os.close(fd)
    os.remove(path)

    # ====================
    # VULN PAGES
    # ====================
    cur_level = ""

    for i, vuln in enumerate(vuln_info, 1):
        # --------------------
        # GENERAL
        # --------------------
        level = vuln.level.lower()

        if level != cur_level:
            document.add_paragraph(
                level.capitalize(), style='OV-H2toc').paragraph_format.page_break_before = True
            cur_level = level
        else:
            document.add_page_break()

        title = "[{}] {}".format(level.upper(), vuln.name)
        document.add_paragraph(title, style='OV-Finding')

        table_vuln = document.add_table(rows=8, cols=3)
        table_vuln.autofit = False

        # COLOR
        # --------------------
        col_cells = table_vuln.columns[0].cells
        col_cells[0].merge(col_cells[7])
        color_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), Config.colors()[vuln.level][1:]))
        col_cells[0]._tc.get_or_add_tcPr().append(color_fill)

        for col_cell in col_cells:
            col_cell.width = Cm(0.42)

        # TABLE HEADERS
        # --------------------
        hdr_cells = table_vuln.columns[1].cells
        hdr_cells[0].paragraphs[0].add_run('Description').bold = True
        hdr_cells[1].paragraphs[0].add_run('Impact').bold = True
        hdr_cells[2].paragraphs[0].add_run('Recommendation').bold = True
        hdr_cells[3].paragraphs[0].add_run('Details').bold = True
        hdr_cells[4].paragraphs[0].add_run('CVSS').bold = True
        hdr_cells[5].paragraphs[0].add_run('CVEs').bold = True
        hdr_cells[6].paragraphs[0].add_run('Family').bold = True
        hdr_cells[7].paragraphs[0].add_run('References').bold = True

        for hdr_cell in hdr_cells:
            hdr_cell.width = Cm(3.58)

        # FIELDS
        # --------------------
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"

        cvss = str(vuln.cvss) if vuln.cvss != -1.0 else "No CVSS"

        txt_cells = table_vuln.columns[2].cells
        txt_cells[0].text = vuln.description
        txt_cells[1].text = vuln.impact
        txt_cells[2].text = vuln.solution
        txt_cells[3].text = vuln.insight
        txt_cells[4].text = cvss
        txt_cells[5].text = cves
        txt_cells[6].text = vuln.family
        txt_cells[7].text = vuln.references

        for txt_cell in txt_cells:
            txt_cell.width = Cm(12.50)

        # VULN HOSTS
        # --------------------
        document.add_paragraph('Vulnerable hosts', style='Heading 4')

        # add coloumn for result per port and resize columns
        table_hosts = document.add_table(cols=5, rows=(len(vuln.hosts) + 1))

        col_cells = table_hosts.columns[1].cells
        for col_cell in col_cells:
            col_cell.width = Cm(3.2)

        col_cells = table_hosts.columns[2].cells
        for col_cell in col_cells:
            col_cell.width = Cm(3.2)

        col_cells = table_hosts.columns[2].cells
        for col_cell in col_cells:
            col_cell.width = Cm(1.6)

        col_cells = table_hosts.columns[3].cells
        for col_cell in col_cells:
            col_cell.width = Cm(1.6)

        col_cells = table_hosts.columns[4].cells
        for col_cell in col_cells:
            col_cell.width = Cm(6.4)

        hdr_cells = table_hosts.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('IP').bold = True
        hdr_cells[1].paragraphs[0].add_run('Host name').bold = True
        hdr_cells[2].paragraphs[0].add_run('Port number').bold = True
        hdr_cells[3].paragraphs[0].add_run('Port protocol').bold = True
        hdr_cells[4].paragraphs[0].add_run('Port result').bold = True

        for j, (host, port) in enumerate(vuln.hosts, 1):
            cells = table_hosts.rows[j].cells
            cells[0].text = host.ip
            cells[1].text = host.host_name if host.host_name else "-"
            if port and port is not None:
                cells[2].text = "-" if port.number == 0 else str(port.number)
                cells[3].text = port.protocol
                cells[4].text = port.result
            else:
                cells[2].text = "No port info"

    document.save(output_file)


def export_to_csv_by_vuln(vuln_info, template=None, output_file='openvas_report.csv'):
    """
    Export vulnerabilities info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")

    vuln_info, _, _, _ = _get_collections(vuln_info)

    with open(output_file, 'w') as csvfile:
        fieldnames = ['hostname', 'ip', 'port', 'protocol',
                      'vulnerability', 'cvss', 'threat', 'family',
                      'description', 'detection', 'insight', 'impact', 'affected', 'solution', 'solution_type',
                      'vuln_id', 'cve', 'references']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for vuln in vuln_info:
            for (host, port) in vuln.hosts:
                rowdata = {
                    'hostname': host.host_name,
                    'ip': host.ip,
                    'port': port.number,
                    'protocol': port.protocol,
                    'vulnerability': vuln.name,
                    'cvss': vuln.cvss,
                    'threat': vuln.level,
                    'family': vuln.family,
                    'description': vuln.description,
                    'detection': vuln.detect,
                    'insight': vuln.insight,
                    'impact': vuln.impact,
                    'affected': vuln.affected,
                    'solution': vuln.solution,
                    'solution_type': vuln.solution_type,
                    'vuln_id': vuln.vuln_id,
                    'cve': ' - '.join(vuln.cves),
                    'references': ' - '.join(vuln.references)
                }
                writer.writerow(rowdata)


def export_to_word_by_host(resulttree: ResultTree, template, output_file='openvas_report.docx'):
    """
    Export vulnerabilities info in a Word file.

    :param resulttree: Vulnerability list info
    :type resulttree: resulttree

    :param output_file: Filename of the Word file
    :type output_file: str

    :param template: Path to Docx template
    :type template: str

    :raises: TypeError
    """

    import matplotlib.pyplot as plt
    import numpy as np
    import tempfile
    import os

    from docx import Document
    from docx.oxml.shared import qn, OxmlElement
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm

    if not isinstance(resulttree, ResultTree):
        raise TypeError("Expected ResultTree, got '{}' instead".format(type(resulttree)))
    else:
        for key in resulttree.keys():
            if not isinstance(resulttree[key], Host):
                raise TypeError("Expected Host, got '{}' instead".format(type(resulttree[key])))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        if not isinstance(template, str):
            raise TypeError("Expected str, got '{}' instead".format(type(template)))
    else:
        # == HAMMER PROGRAMMING (beat it into submission) ==
        # I had to use pkg_resources because I couldn't find this template any other way.
        import pkg_resources
        template = pkg_resources.resource_filename('openvasreporting', 'src/openvas-template.docx')

    # mpl_logger = logging.getLogger('matplotlib')
    # mpl_logger.setLevel(logging.NOTSET)

    # ====================
    # DOCUMENT PROPERTIES
    # ====================
    document = Document(template)

    doc_prop = document.core_properties
    doc_prop.title = "OpenVAS Report"
    doc_prop.category = "Report"

    temp_resulttree = resulttree.sorted_keys_by_rank()
    for i, key in enumerate(temp_resulttree, 1):

        # this host has any vulnerability whose cvss severity >= min_level?
        if len(resulttree[key].vuln_list) == 0:
            continue

        # --------------------
        # HOST SUMMARY TABLE
        # --------------------
        document.add_paragraph('Summary {}'.format(resulttree[key].ip + ' - ' + resulttree[key].host_name))


        # table_host_summary = document.add_table(rows=1, cols=8)
        #
        # hdr_cells = table_host_summary.rows[0].cells
        # hdr_cells[0].text = "CVSS"
        # hdr_cells[1].text = "Name"
        # hdr_cells[2].text = "oid"
        # hdr_cells[3].text = "Port"
        # hdr_cells[4].text = "Family"
        # hdr_cells[5].text = "Description"
        # hdr_cells[6].text = "Recommendation"
        # hdr_cells[7].text = "Type of fix"

        # --------------------
        # HOST VULN LIST
        # --------------------
        for j, vuln in enumerate(resulttree[key].vuln_list):
            port = vuln.hosts[0][1]
            if port is None or port.number == 0:
                portnum = 'general'
            else:
                portnum = str(port.number)
            section = 'CVSS:\n{:.2f} ({})\n\nName:\n{}\n\noid:\n{}\n\nPort:\n{}\n\nFamily:\n{}\n\nDescription:\n{}\n\nRecommendation:\n{}\n\nType of fix:\n{}\n\n'.format(
                vuln.cvss,
                vuln.level,
                vuln.name,
                vuln.vuln_id,
                '{}/{}'.format(portnum, port.protocol),
                vuln.family,
                vuln.description.replace('\n', ' '),
                vuln.solution.replace('\n', ' '),
                vuln.solution_type
            )
            document.add_paragraph(section)
            # vuln_row_cells = table_host_summary.add_row().cells
            # vuln_row_cells[0].text = '{:.2f} ({})'.format(vuln.cvss, vuln.level)
            # vuln_row_cells[1].text = vuln.name
            # vuln_row_cells[2].text = vuln.vuln_id
            # port = vuln.hosts[0][1]
            # if port is None or port.number == 0:
            #     portnum = 'general'
            # else:
            #     portnum = str(port.number)
            # vuln_row_cells[3].text = '{}/{}'.format(portnum, port.protocol)
            # vuln_row_cells[4].text = vuln.family
            # vuln_row_cells[5].text = vuln.description.replace('\n', ' ')
            # vuln_row_cells[6].text = vuln.solution.replace('\n', ' ')
            # vuln_row_cells[7].text = vuln.solution_type
        document.add_page_break()

    document.save(output_file)


def export_to_excel_by_host(resulttree: ResultTree, template=None, output_file='openvas_report.xlsx'):
    """
    Export vulnerabilities info in an Excel file.

    :param resulttree: Vulnerability list info
    :type resulttree: resulttree
    :param template: Not supported in xlsx-output
    :type template: NoneType

    :param output_file: Filename of the Excel file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import xlsxwriter

    if not isinstance(resulttree, ResultTree):
        raise TypeError("Expected ResultTree, got '{}' instead".format(type(resulttree)))
    else:
        for key in resulttree.keys():
            if not isinstance(resulttree[key], Host):
                raise TypeError("Expected Host, got '{}' instead".format(type(resulttree[key])))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in XSLX-output.")

    # ====================
    # FUNCTIONS
    # ====================
    def __row_height(text, width):
        return (max((len(text) // width), text.count('\n')) + 1) * 15

    workbook = xlsxwriter.Workbook(output_file)

    workbook.set_properties({
        'title': output_file,
        'subject': 'OpenVAS report',
        'author': 'TheGroundZero, ecgf(IcatuHolding)',
        'category': 'report',
        'keywords': 'OpenVAS, report',
        'comments': 'TheGroundZero (https://github.com/TheGroundZero)'})

    # ====================
    # FORMATTING
    # ====================
    workbook.formats[0].set_font_name('Tahoma')

    format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'center', 'valign': 'vcenter', 'border': 1})
    format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                               'font_color': 'white', 'bold': True,
                                               'align': 'center', 'valign': 'vcenter',
                                               'border': 1,
                                               'bg_color': Config.colors()['blue']})
    format_table_left_item = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'left', 'valign': 'vcenter', 'border': 1})
    format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                              'align': 'left', 'valign': 'top',
                                              'border': 1, 'text_wrap': 1})
    format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top'})
    format_align_left = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'left', 'valign': 'top'})
    format_align_right = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'right', 'valign': 'top'})
    format_align_border_left = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'left', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_align_border_right = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'right', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_number_border_right = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'right', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_number_border_right.num_format = '#.00'
    format_toc = {
        'critical': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                         'align': 'center', 'valign': 'top',
                                         'border': 1,
                                         'bg_color': Config.colors()['critical']}),
        'high': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['high']}),
        'medium': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                       'align': 'center', 'valign': 'top',
                                       'border': 1, 'bg_color': Config.colors()['medium']}),
        'low': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                    'align': 'center', 'valign': 'top',
                                    'border': 1, 'bg_color': Config.colors()['low']}),
        'none': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['none']})
    }

    # ====================
    # SUMMARY SHEET
    # ====================
    sheet_name = "Summary"
    ws_sum = workbook.add_worksheet(sheet_name)
    ws_sum.set_tab_color(Config.colors()['blue'])

    ws_sum.set_column("A:A", 3, format_align_center)
    ws_sum.set_column("B:B", 8, format_align_left)
    ws_sum.set_column("C:C", 30, format_align_left)
    ws_sum.set_column("D:D", 15, format_align_right) # critical
    ws_sum.set_column("E:E", 8, format_align_right) # high
    ws_sum.set_column("F:F", 8, format_align_right) # medium
    ws_sum.set_column("G:G", 8, format_align_right) # low
    ws_sum.set_column("H:H", 8, format_align_right) # none
    ws_sum.set_column("I:I", 8, format_align_right) # total
    ws_sum.set_column("J:J", 8, format_align_right) # severity
    ws_sum.set_column("K:K", 7, format_align_center)

    # ---------------------
    # MAX 10 HOSTS 
    # ---------------------
    if len(resulttree) < 10:
        max_hosts = len(resulttree)
    else:
        max_hosts = 10

    # --------------------------
    # HOST SUM SEVERITY SUMMARY
    # --------------------------
    ws_sum.merge_range("B2:J2", "Hosts Ranking", format_sheet_title_content)
    ws_sum.write("B3", "#", format_table_titles)
    ws_sum.write("C3", "Hostname", format_table_titles)
    ws_sum.write("D3", "IP", format_table_titles)
    ws_sum.write("E3", "critical", format_table_titles)
    ws_sum.write("F3", "high", format_table_titles)
    ws_sum.write("G3", "medium", format_table_titles)
    ws_sum.write("H3", "low", format_table_titles)
    ws_sum.write("I3", "total", format_table_titles)
    ws_sum.write("J3", "severity", format_table_titles)
    
    temp_resulttree = resulttree.sorted_keys_by_rank()
    
    for i, key in enumerate(temp_resulttree[:max_hosts], 4):
        ws_sum.write("B{}".format(i), i-3, format_table_left_item)
        ws_sum.write("C{}".format(i), resulttree[key].host_name, format_table_left_item)
        ws_sum.write("D{}".format(i), resulttree[key].ip, format_table_left_item)
        ws_sum.write("E{}".format(i), resulttree[key].nv['critical'], format_align_border_right)
        ws_sum.write("F{}".format(i), resulttree[key].nv['high'], format_align_border_right)
        ws_sum.write("G{}".format(i), resulttree[key].nv['medium'], format_align_border_right)
        ws_sum.write("H{}".format(i), resulttree[key].nv['low'], format_align_border_right)
        ws_sum.write("I{}".format(i), resulttree[key].nv_total(), format_align_border_right)
        ws_sum.write("J{}".format(i), resulttree[key].higher_cvss, 
                     format_toc[Config.cvss_level(resulttree[key].higher_cvss)])

    # --------------------
    # CHART
    # --------------------
    chart_sumcvss_summary = workbook.add_chart({'type': 'column'})
    chart_sumcvss_summary.add_series({
        'name': 'critical',
        'categories': '={}!D4:D{}'.format(sheet_name, max_hosts + 3),
        'values': '={}!E4:E{}'.format(sheet_name, max_hosts + 3),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma', 'size': 8}},
        'fill': { 'width': 8, 'color': Config.colors()['critical']},
        'border': { 'color': Config.colors()['blue']},
    })
    chart_sumcvss_summary.add_series({
        'name': 'high',
        'values': '={}!F4:F{}'.format(sheet_name, max_hosts + 3),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma', 'size': 8}},
        'fill': { 'width': 8, 'color': Config.colors()['high']},
        'border': { 'color': Config.colors()['blue']},
    })
    chart_sumcvss_summary.add_series({
        'name': 'medium',
        'values': '={}!G4:G{}'.format(sheet_name, max_hosts + 3),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma', 'size': 8}},
        'fill': { 'width': 8, 'color': Config.colors()['medium']},
        'border': { 'color': Config.colors()['blue']},
    })
    
    #chart_sumcvss_summary.add_series({
        #'name': 'Hosts Ranking',
        #'categories': '={}!D4:D{}'.format(sheet_name, max_hosts + 3),
        #'values': '={}!E4:G{}'.format(sheet_name, max_hosts + 3),
        #'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma', 'size': 8}},
        #'line': { 'width': 8, 'color': Config.colors()['blue']},
        #'border': { 'color': Config.colors()['blue']},
    #})
    chart_sumcvss_summary.set_title({'name': 'Hosts by CVSS', 'overlay': False, 'font': {'name': 'Tahoma'}})
    chart_sumcvss_summary.set_size({'width': 750, 'height': 350})
    chart_sumcvss_summary.set_legend({'position': 'left', 'font': {'name': 'Tahoma'}})
    chart_sumcvss_summary.set_x_axis({'label_position': 'bottom',
                                      'num_font': {'name': 'Tahoma', 'size': 8}
                                    })
    ws_sum.insert_chart("B15", chart_sumcvss_summary)

    # ====================
    # TABLE OF CONTENTS
    # ====================
    sheet_name = "TOC"
    ws_toc = workbook.add_worksheet(sheet_name)
    ws_toc.set_tab_color(Config.colors()['blue'])

    ws_toc.set_column("A:A", 3, format_align_center)
    ws_toc.set_column("B:B", 8, format_align_left)
    ws_toc.set_column("C:C", 30, format_align_left)
    ws_toc.set_column("D:D", 15, format_align_right) # critical
    ws_toc.set_column("E:E", 8, format_align_right) # high
    ws_toc.set_column("F:F", 8, format_align_right) # medium
    ws_toc.set_column("G:G", 8, format_align_right) # low
    ws_toc.set_column("H:H", 8, format_align_right) # none
    ws_toc.set_column("I:I", 8, format_align_right) # total
    ws_toc.set_column("J:J", 8, format_align_right) # severity
    ws_toc.set_column("K:K", 7, format_align_center)

    # --------------------------
    # HOST SUM SEVERITY SUMMARY
    # --------------------------
    ws_toc.merge_range("B2:J2", "Hosts Ranking", format_sheet_title_content)
    ws_toc.write("B3", "#", format_table_titles)
    ws_toc.write("C3", "Hostname", format_table_titles)
    ws_toc.write("D3", "IP", format_table_titles)
    ws_toc.write("E3", "critical", format_table_titles)
    ws_toc.write("F3", "high", format_table_titles)
    ws_toc.write("G3", "medium", format_table_titles)
    ws_toc.write("H3", "low", format_table_titles)
    ws_toc.write("I3", "total", format_table_titles)
    ws_toc.write("J3", "severity", format_table_titles)
    
    # ====================
    # HOST SHEETS
    # ====================
    for i, key in enumerate(temp_resulttree, 1):
 
        # this host has any vulnerability whose cvss severity >= min_level?
        if len(resulttree[key].vuln_list) == 0:
            continue

        name = "{:03X} - {}".format(i, resulttree[key].ip)
        ws_host = workbook.add_worksheet(name)
        ws_host.set_tab_color(Config.cvss_color(resulttree[key].higher_cvss))
        ws_host.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), i + 3), format_align_center,
                          string="<< TOC")

        # --------------------
        # TABLE OF CONTENTS
        # --------------------
        ws_toc.write("B{}".format(i + 3), "{:03X}".format(i), format_table_cells)
        ws_toc.write_url("C{}".format(i + 3), "internal:'{}'!A1".format(name), format_table_cells, 
                         string=resulttree[key].host_name)
        ws_toc.write("D{}".format(i+3), resulttree[key].ip, format_align_border_left)
        ws_toc.write("E{}".format(i+3), resulttree[key].nv['critical'], format_align_border_right)
        ws_toc.write("F{}".format(i+3), resulttree[key].nv['high'], format_align_border_right)
        ws_toc.write("G{}".format(i+3), resulttree[key].nv['medium'], format_align_border_right)
        ws_toc.write("H{}".format(i+3), resulttree[key].nv['low'], format_align_border_right)
        ws_toc.write("I{}".format(i+3), resulttree[key].nv_total(), format_align_border_right)
        ws_toc.write("J{}".format(i+3), resulttree[key].higher_cvss, 
                     format_toc[Config.cvss_level(resulttree[key].higher_cvss)])
        ws_toc.set_row(i + 3, __row_height(name, 150), None)

        # --------------------
        # HOST VULN LIST
        # --------------------
        ws_host.set_column("A:A", 7, format_align_center)
        ws_host.set_column("B:B", 12, format_align_center) # cvss - (level)
        ws_host.set_column("C:C", 22, format_align_center) # name
        ws_host.set_column("D:D", 22, format_align_center) # oid
        ws_host.set_column("E:E", 10, format_align_center) # port.port/port.num
        ws_host.set_column("F:F", 10, format_align_center) # family
        ws_host.set_column("G:G", 22, format_align_center) # description
        ws_host.set_column("H:H", 22, format_align_center) # recomendation (solution)
        ws_host.set_column("I:I", 12, format_align_center) # recomendation type (solution_type)
        ws_host.set_column("J:J", 7, format_align_center)
        
        ws_host.merge_range("B2:I2", resulttree[key].ip + ' - ' + resulttree[key].host_name, format_sheet_title_content)
        ws_host.write('B3', "CVSS", format_table_titles)
        ws_host.write('C3', "Name", format_table_titles)
        ws_host.write('D3', "oid", format_table_titles)
        ws_host.write('E3', "Port", format_table_titles)
        ws_host.write('F3', "Family", format_table_titles)
        ws_host.write('G3', "Description", format_table_titles)
        ws_host.write('H3', "Recomendation", format_table_titles)
        ws_host.write('I3', "Type of fix", format_table_titles)


        for j, vuln in enumerate(resulttree[key].vuln_list, 4):
            ws_host.write('B{}'.format(j), "{:.2f} ({})".format(vuln.cvss, vuln.level),
                          format_toc[vuln.level])
            ws_host.write('C{}'.format(j), vuln.name, format_align_border_left)
            ws_host.write('D{}'.format(j), vuln.vuln_id, format_align_border_left)
            port = vuln.hosts[0][1]
            if port is None or port.number == 0:
                portnum = 'general' 
            else: 
                portnum = str(port.number)
            ws_host.write('E{}'.format(j), portnum + '/' + port.protocol, format_align_border_left)
            ws_host.write('F{}'.format(j), vuln.family, format_align_border_left)
            ws_host.write('G{}'.format(j), vuln.description.replace('\n', ' '), format_align_border_left)
            ws_host.write('H{}'.format(j), vuln.solution.replace('\n', ' '), format_align_border_left)
            ws_host.write('I{}'.format(j), vuln.solution_type, format_align_border_left)
            max_len = max(len(vuln.name), len(vuln.description), len(vuln.solution))
            ws_host.set_row(j-1, (int(max_len/30)+1)*15)
        
    workbook.close()


def export_to_csv_by_host(resulttree, template=None, output_file='openvas_report.csv'):
    """
    Export vulnerabilities info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv
    
    if not isinstance(resulttree, ResultTree):
        raise TypeError("Expected ResultTree, got '{}' instead".format(type(resulttree)))
    else:
        for x in resulttree:
            if not isinstance(x, Host):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")
  
    sortedresults = resulttree.sortedbysumcvss()

    with open(output_file, 'w') as csvfile:
        fieldnames = ['hostname', 'ip', 'port', 'protocol',
                      'vulnerability', 'cvss', 'threat', 'family',
                      'description', 'detection', 'insight', 'impact', 'affected', 'solution', 'solution_type',
                      'vuln_id', 'cve', 'references']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for key in sortedresults:
            for vuln in resulttree[key].vuln_list:
                rowdata = {
                    'hostname': resulttree[key].host_name,
                    'ip': resulttree[key].ip,
                    'port': vuln.port.number,
                    'protocol': vuln.port.protocol,
                    'vulnerability': vuln.name,
                    'cvss': vuln.cvss,
                    'threat': vuln.level,
                    'family': vuln.family,
                    'description': vuln.description,
                    'detection': vuln.detect,
                    'insight': vuln.insight,
                    'impact': vuln.impact,
                    'affected': vuln.affected,
                    'solution': vuln.solution,
                    'solution_type': vuln.solution_type,
                    'vuln_id': vuln.vuln_id,
                    'cve': ' - '.join(vuln.cves),
                    'references': ' - '.join(vuln.references) if isinstance(vuln.references, list) else vuln.references
                }
                writer.writerow(rowdata)

def export_summary_to_csv(
        vuln_info,
        template=None,
        output_file='openvas_summary_report.csv'
    ):
    """
    Export summary info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")

    vuln_info, vuln_levels, vuln_host_by_level, _ = _get_collections(vuln_info)

    with open(output_file, 'w') as csvfile:
        fieldnames = ['level', 'count', 'host_count']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for i, level in enumerate(Config.levels().values(), 4):
            rowdata = {
                'level': level,
                'count': vuln_levels[level],
                'host_count': vuln_host_by_level[level]
            }
            writer.writerow(rowdata)
